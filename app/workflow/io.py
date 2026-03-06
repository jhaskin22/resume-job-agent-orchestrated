import logging
import re
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

class ResumeParsingError(RuntimeError):
    pass


def parse_resume_content(filename: str, payload: bytes) -> str:
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return _parse_pdf(payload)
    if extension == ".docx":
        return _parse_docx(payload)

    raise ResumeParsingError("Unsupported resume format. Accepted types: PDF, DOCX")


def _parse_pdf(payload: bytes) -> str:
    reader = PdfReader(BytesIO(payload))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(part.strip() for part in pages if part.strip())


def _parse_docx(payload: bytes) -> str:
    doc = Document(BytesIO(payload))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


def safe_stem(raw_filename: str) -> str:
    base = Path(raw_filename).stem.lower()
    base = re.sub(r"[^a-z0-9]+", "-", base).strip("-")
    return base or "resume"


def write_resume_docx(content: str, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for block in content.split("\n"):
        clean = block.strip()
        if clean:
            doc.add_paragraph(clean)
    doc.save(str(output_path))


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return "\n".join(paragraphs)


def persist_uploaded_resume(payload: bytes, filename: str, upload_dir: Path) -> Path:
    upload_dir.mkdir(parents=True, exist_ok=True)
    ext = Path(filename).suffix.lower() or ".docx"
    stem = safe_stem(filename)
    stored_name = f"{stem}-{uuid4().hex[:10]}{ext}"
    path = upload_dir / stored_name
    path.write_bytes(payload)
    return path


def rewrite_resume_docx(
    payload: bytes,
    output_path: Path,
    emphasis_keywords: list[str],
    max_rewrites: int = 2,
) -> dict[str, int]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(BytesIO(payload))

    original_paragraphs = len(doc.paragraphs)
    bullet_indexes = _bullet_paragraph_indexes(doc)
    rewrite_indexes = bullet_indexes or _experience_paragraph_indexes(doc)
    modified_bullets = 0

    for paragraph_index in rewrite_indexes[: min(max_rewrites, 2)]:
        paragraph = doc.paragraphs[paragraph_index]
        original_text = paragraph.text
        rewritten = _rewrite_bullet_text(original_text, emphasis_keywords, modified_bullets)
        if rewritten != original_text:
            paragraph.text = rewritten
            modified_bullets += 1

    doc.save(str(output_path))

    logger.info(
        "resume_rewrite paragraphs=%s bullets=%s modified=%s keywords=%s",
        original_paragraphs,
        len(rewrite_indexes),
        modified_bullets,
        len(emphasis_keywords),
    )

    return {
        "original_paragraphs": original_paragraphs,
        "output_paragraphs": len(doc.paragraphs),
        "bullet_count": len(rewrite_indexes),
        "modified_bullets": modified_bullets,
    }


def _bullet_paragraph_indexes(doc: Document) -> list[int]:
    indexes: list[int] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if (
            "list bullet" in style_name
            or text.startswith(("-", "•", "*"))
            or text[:2].isdigit() and text[1:2] == "."
        ):
            indexes.append(idx)
    return indexes


def _experience_paragraph_indexes(doc: Document) -> list[int]:
    indexes: list[int] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if len(text.split()) < 5:
            continue
        lowered = text.lower()
        if re.search(r"\b(summary|experience|skills|education|projects)\b", lowered):
            continue
        if re.search(r"\b(led|built|developed|designed|managed|improved|implemented)\b", lowered):
            indexes.append(idx)
    return indexes


def _rewrite_bullet_text(text: str, keywords: list[str], variant_index: int) -> str:
    stripped = text.strip()
    prefix = ""
    body = stripped
    if stripped.startswith(("•", "-", "*")):
        prefix = f"{stripped[0]} "
        body = stripped[1:].strip()

    verb_map = {
        "design": "designed",
        "operate": "operated",
        "build": "built",
        "develop": "developed",
        "lead": "led",
        "manage": "managed",
        "improve": "improved",
    }
    rewritten = body
    for base, replacement in verb_map.items():
        rewritten = re.sub(
            rf"\b{base}(ed|ing|s)?\b",
            replacement,
            rewritten,
            count=1,
            flags=re.IGNORECASE,
        )

    if keywords:
        keyword = _keyword_for_sentence(keywords, variant_index)
        if keyword and keyword.lower() not in rewritten.lower():
            rewritten = _apply_keyword_context(rewritten, keyword)

    rewritten = rewritten[0].upper() + rewritten[1:] if rewritten else rewritten
    if not rewritten.endswith("."):
        rewritten = f"{rewritten}."
    return f"{prefix}{rewritten}".strip()


def _keyword_for_sentence(keywords: list[str], variant_index: int) -> str:
    banned = {"engineer", "aligned", "role", "job", "team", "company", "experience"}
    filtered = [
        keyword.strip()
        for keyword in keywords
        if keyword.strip() and keyword.strip().lower() not in banned and len(keyword.strip()) >= 3
    ]
    if not filtered:
        return ""
    return filtered[variant_index % len(filtered)]


def _apply_keyword_context(sentence: str, keyword: str) -> str:
    patterns = (
        (r"\bbackend services\b", f"{keyword} backend services"),
        (r"\bsoftware systems\b", f"{keyword} software systems"),
        (r"\bapi services\b", f"{keyword} API services"),
        (r"\bworkflows\b", f"{keyword} workflows"),
    )
    for pattern, replacement in patterns:
        if re.search(pattern, sentence, flags=re.IGNORECASE):
            return re.sub(pattern, replacement, sentence, count=1, flags=re.IGNORECASE)
    verb_match = re.match(
        r"^(Built|Developed|Designed|Managed|Led|Improved|Implemented)\b",
        sentence,
        flags=re.IGNORECASE,
    )
    if verb_match:
        verb = verb_match.group(1)
        return re.sub(
            r"^(Built|Developed|Designed|Managed|Led|Improved|Implemented)\b",
            f"{verb} scalable {keyword}",
            sentence,
            count=1,
            flags=re.IGNORECASE,
        )
    return f"{keyword} {sentence}".strip()
