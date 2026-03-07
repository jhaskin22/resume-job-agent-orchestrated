import asyncio
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient

from app.core.config import settings
from app.main import app


@pytest.fixture
def anyio_backend() -> str:
    return "asyncio"


def _build_docx_payload() -> bytes:
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Senior Python engineer with AI workflow and API development experience.")
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(
        "• Design and operate high-reliability, low-latency software systems for telemetry."
    )
    doc.add_paragraph(
        "• Build API services and automation tooling to improve testing and deployment confidence."
    )
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, FastAPI, Docker, Kubernetes, AWS, LangGraph")
    doc.save(buffer)
    return buffer.getvalue()


@pytest.mark.anyio
async def test_run_workflow_with_docx() -> None:
    upload_dir = Path(settings.uploaded_resume_dir)
    before_uploads = {path.name for path in upload_dir.glob("*")} if upload_dir.exists() else set()

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        response = await client.post(
            "/api/workflow/run",
            files={
                "resume": (
                    "candidate.docx",
                    _build_docx_payload(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload["tiles"]
    assert payload["diagnostics"]["verification"]["job_parsing"]["ok"] is True
    assert payload["diagnostics"]["verification"]["tile_construction"]["ok"] is True
    first = payload["tiles"][0]
    assert first["company"]
    assert first["generated_resume_link"] == ""
    assert "verification" in payload["diagnostics"]

    after_uploads = {path.name for path in upload_dir.glob("*")} if upload_dir.exists() else set()
    assert len(after_uploads - before_uploads) >= 1


@pytest.mark.anyio
async def test_generated_resume_download_on_demand() -> None:
    original_payload = _build_docx_payload()
    original_doc = Document(BytesIO(original_payload))
    original_paragraph_count = len(original_doc.paragraphs)

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        start_response = await client.post(
            "/api/workflow/runs",
            files={
                "resume": (
                    "candidate.docx",
                    original_payload,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert start_response.status_code == 200
        run_id = start_response.json()["run_id"]
        payload = {}
        for _ in range(300):
            status_response = await client.get(f"/api/workflow/runs/{run_id}")
            assert status_response.status_code == 200
            payload = status_response.json()
            if payload["status"] == "failed":
                break
            if len(payload.get("tiles", [])) >= 1:
                break
            await asyncio.sleep(0.2)
        assert payload.get("status") != "failed"
        assert len(payload["tiles"]) >= 1
        first_job_link = payload["tiles"][0]["job_link"]

        gen_response = await client.post(
            f"/api/workflow/runs/{run_id}/resume",
            json={"job_link": first_job_link},
        )
        assert gen_response.status_code == 200
        link = gen_response.json()["generated_resume_link"]

        download_response = await client.get(link)

    assert download_response.status_code == 200
    assert download_response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(download_response.content) > 100

    generated_doc = Document(BytesIO(download_response.content))
    assert len(generated_doc.paragraphs) == original_paragraph_count
